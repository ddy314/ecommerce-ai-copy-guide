from __future__ import annotations

import json
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "submission-output"
ASSETS = OUT / "assets"
TODAY = date(2026, 7, 15)
PROJECT = "电商 AI 商品文案生成与智能导购助手"
DOC_CODE = "EACG-2026"

NAVY = "17324D"
BLUE = "2563EB"
CYAN = "0891B2"
MINT = "0F766E"
INK = "172033"
MUTED = "64748B"
LIGHT = "EEF4FA"
PALE = "F8FAFC"
LINE = "CBD5E1"
WHITE = "FFFFFF"
RED = "B91C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color=INK, font="Microsoft YaHei") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run(run, 9, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[list], widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(text)
        set_cell_shading(cell, NAVY)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run(run, font_size, True, WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.text = "" if value is None else str(value)
            if ridx % 2 == 1:
                set_cell_shading(cell, PALE)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                set_run(run, font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc: Document, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_callout(doc: Document, label: str, text: str, color=MINT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "ECFDF5" if color == MINT else LIGHT)
    set_cell_border(cell, color, "8")
    set_cell_margins(cell, 140, 170, 140, 170)
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    set_run(r, 10, True, color)
    r = p.add_run(text)
    set_run(r, 10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, path: Path, caption: str, width=6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    for run in c.runs:
        set_run(run, 9, color=MUTED)


def new_doc(title: str, doc_no: str, version="1.0", subtitle="课程项目交付文档") -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.25)
    sec.right_margin = Cm(2.05)
    sec.header_distance = Cm(0.9)
    sec.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 8),
        ("Heading 1", 16, NAVY, 15, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, CYAN, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = name != "Title" or True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header
    hp = header.paragraphs[0]
    hp.text = f"{PROJECT}  |  {title}"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        set_run(run, 8.5, color=MUTED)
    add_page_field(sec.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("E-COMMERCE  ·  AI  ·  ASSISTANT")
    set_run(r, 10, True, CYAN)
    p = doc.add_paragraph()
    p.style = "Title"
    p.add_run(title)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(PROJECT)
    set_run(r, 17, True, BLUE)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    set_run(r, 11, color=MUTED)

    add_table(
        doc,
        ["文档编号", "版本", "状态", "日期"],
        [[doc_no, version, "提交版", TODAY.isoformat()]],
        [4.4, 3.0, 3.0, 4.0],
        9.5,
    )
    doc.add_paragraph()
    add_callout(doc, "文档说明", "本文档依据当前仓库实现编制；人员姓名与签字栏保留为可填写项，不虚构项目成员。", BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目组 · 2026")
    set_run(r, 11, True, NAVY)
    doc.add_page_break()
    return doc


def add_front_matter(doc: Document, title: str, revisions=None, contents=None) -> None:
    doc.add_heading("文档信息", level=1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["文档标题", title],
            ["适用项目", PROJECT],
            ["编制角色", "项目组（姓名待补充）"],
            ["评审角色", "指导教师 / 项目负责人"],
            ["密级", "课程内部"],
        ],
        [3.6, 11.0],
    )
    doc.add_heading("修订记录", level=1)
    revisions = revisions or [[TODAY.isoformat(), "1.0", "按当前项目实现创建提交版", "项目组"]]
    add_table(doc, ["日期", "版本", "说明", "作者"], revisions, [3.2, 2.4, 7.0, 2.4])
    if contents:
        doc.add_heading("目录", level=1)
        for idx, item in enumerate(contents, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(f"{idx:02d}  {item}")
            set_run(r, 10.5, idx <= 3, NAVY if idx <= 3 else INK)
    doc.add_page_break()


def save_doc(doc: Document, name: str) -> Path:
    path = OUT / name
    props = doc.core_properties
    props.title = name.removesuffix(".docx")
    props.subject = PROJECT
    props.author = "项目组"
    props.keywords = "电商, AI, Flask, Vue, 项目交付"
    props.comments = "Generated from the current repository implementation."
    doc.save(path)
    return path


def make_diagrams() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "system-architecture": r'''
digraph G {
  graph [rankdir=LR, bgcolor="transparent", pad=0.2, nodesep=0.45, ranksep=0.8, splines=ortho];
  node [shape=box, style="rounded,filled", fontname="Microsoft YaHei", fontsize=12, margin="0.18,0.12", color="#CBD5E1", penwidth=1.2];
  edge [color="#64748B", penwidth=1.4, arrowsize=0.75];
  browser [label="浏览器\n用户端 / 商家端", fillcolor="#E0F2FE", color="#0891B2"];
  nginx [label="Nginx :8080\n静态资源 + 反向代理", fillcolor="#ECFDF5", color="#0F766E"];
  flask [label="Flask / Gunicorn :8000\nREST API · JWT · CORS", fillcolor="#DBEAFE", color="#2563EB"];
  db [label="PostgreSQL 16\n本地可用 SQLite", shape=cylinder, fillcolor="#F8FAFC", color="#475569"];
  redis [label="Redis 7\n可降级缓存", shape=cylinder, fillcolor="#FEF3C7", color="#D97706"];
  ai [label="OpenAI 兼容接口\n或确定性 Mock AI", fillcolor="#F3E8FF", color="#7C3AED"];
  crawler [label="京东 / 苏宁\n商品与评论采集", fillcolor="#FFF1F2", color="#E11D48"];
  browser -> nginx -> flask;
  flask -> db;
  flask -> redis;
  flask -> ai;
  crawler -> flask;
}
''',
        "data-relationships": r'''
digraph G {
  graph [rankdir=LR, bgcolor="transparent", pad=0.2, nodesep=0.35, ranksep=0.8, splines=ortho];
  node [shape=record, style="rounded,filled", fontname="Microsoft YaHei", fontsize=10, fillcolor="#F8FAFC", color="#94A3B8"];
  edge [color="#64748B", arrowsize=0.7];
  users [label="{users|身份 · 资料 · 角色}", fillcolor="#DBEAFE", color="#2563EB"];
  products [label="{products|商品 · 价格 · 媒体}", fillcolor="#ECFDF5", color="#0F766E"];
  reviews [label="{reviews|评分 · 文本 · 媒体}"];
  cart [label="{cart_items|用户 · 商品 · 数量}"];
  orders [label="{orders|状态 · 金额 · 地址快照}", fillcolor="#FEF3C7", color="#D97706"];
  items [label="{order_items|商品快照 · 数量}"];
  knowledge [label="{knowledge_entries|规格 · FAQ · 售后}", fillcolor="#F3E8FF", color="#7C3AED"];
  qa [label="{qa_records|问题 · 回答 · 来源}"];
  cs [label="{customer_service_messages|会话 · 已读状态}"];
  tasks [label="{generation_tasks|输入 · 输出 · 状态}"];
  users -> cart;
  products -> cart;
  users -> orders;
  orders -> items;
  products -> items;
  products -> reviews;
  products -> knowledge;
  users -> qa;
  products -> qa;
  users -> cs;
  products -> cs;
  products -> tasks;
}
''',
        "ai-flow": r'''
digraph G {
  graph [rankdir=LR, bgcolor="transparent", pad=0.2, nodesep=0.4, ranksep=0.7, splines=ortho];
  node [shape=box, style="rounded,filled", fontname="Microsoft YaHei", fontsize=11, fillcolor="#F8FAFC", color="#94A3B8"];
  edge [color="#2563EB", penwidth=1.4, arrowsize=0.75];
  input [label="结构化输入\n商品 / 需求 / 评论"];
  validate [label="Pydantic 校验", fillcolor="#DBEAFE", color="#2563EB"];
  context [label="商品数据 + 知识库\n构造上下文", fillcolor="#ECFDF5", color="#0F766E"];
  provider [label="AI Provider\nOpenAI / Mock", fillcolor="#F3E8FF", color="#7C3AED"];
  output [label="结构化输出\n文案 / 推荐 / 洞察"];
  persist [label="任务与问答记录\n写入数据库", fillcolor="#FEF3C7", color="#D97706"];
  input -> validate -> context -> provider -> output -> persist;
}
''',
    }
    for name, source in diagrams.items():
        dot_path = ASSETS / f"{name}.dot"
        dot_path.write_text(source, encoding="utf-8")
        subprocess.run(["dot", "-Tpng", "-Gdpi=180", str(dot_path), "-o", str(ASSETS / f"{name}.png")], check=True)


def requirements_doc() -> Path:
    title = "软件需求规约"
    doc = new_doc(title, f"{DOC_CODE}-SRS-001")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["引言", "软件总体概述", "用户与角色", "功能需求", "外部接口需求", "非功能需求", "数据与合规要求", "验收标准", "需求追踪矩阵"])
    doc.add_heading("1 引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    doc.add_paragraph("本文档定义系统边界、用户角色、核心功能、质量属性与验收口径，作为设计、开发、测试、答辩和后续迭代的共同基线。")
    doc.add_heading("1.2 项目范围", level=2)
    doc.add_paragraph("系统面向电商运营人员和普通购物用户，覆盖商品采集、商品管理、AI 文案、智能导购、评论洞察、直播脚本、购物交易、售后、知识库与客服等场景。课程演示模式允许在无外部 AI 密钥、无 Redis、无 PostgreSQL 的条件下使用 Mock AI 与 SQLite 完成核心流程。")
    doc.add_heading("1.3 术语", level=2)
    add_table(doc, ["术语", "含义"], [["RAG", "检索增强生成；先从商品/知识库取得上下文，再生成回答"], ["Mock AI", "无需外部密钥、返回可重复结果的演示提供者"], ["JWT", "用于 API 身份认证的 Bearer 令牌"], ["商家端", "运营、商品、订单、用户、知识库与客服管理界面"], ["用户端", "商品浏览、购物车、订单、售后、评价、问答与个人资料界面"]], [3.0, 11.6])

    doc.add_heading("2 软件总体概述", level=1)
    add_figure(doc, ASSETS / "system-architecture.png", "图 1  系统边界与外部依赖")
    doc.add_paragraph("系统采用前后端分离的单页应用架构。Vue 3 前端通过同源 /api 访问 Flask 服务；生产环境由 Nginx 提供静态资源并反向代理，后端使用 SQLAlchemy 访问数据库，Redis 作为可降级缓存，AI Provider 在 OpenAI 兼容接口与 Mock 实现之间切换。")
    doc.add_heading("2.1 产品目标", level=2)
    for item in ["降低电商运营人员从商品信息到营销内容的整理成本。", "让用户通过自然语言获得可解释的商品推荐和问答。", "将评论、订单、客服和知识库数据纳入一个可演示的业务闭环。", "保持离线可演示、容器可部署、接口可扩展。"]:
        add_bullet(doc, item)
    doc.add_heading("2.2 用户与角色", level=2)
    add_table(doc, ["角色", "主要目标", "关键权限"], [["游客", "了解系统并注册", "查看登录/注册入口"], ["普通用户", "浏览、购买、评价与咨询", "商品、购物车、订单、地址、收藏、问答、客服"], ["商家管理员", "运营商品和服务用户", "看板、商品、订单、用户、知识库、问答统计、客服、采集任务"], ["系统维护者", "部署与排障", "配置环境变量、迁移、日志与健康检查"]], [2.8, 5.2, 6.6])

    doc.add_heading("3 功能需求", level=1)
    functional = [
        ("FR-01", "认证与资料", "注册、登录、身份匹配校验、找回/重置密码、资料与头像更新、退出登录；JWT 有效期 24 小时。"),
        ("FR-02", "商品与评论", "分页、分类、关键词查询商品，查看详情与评论；商家可新增、编辑、删除、批量上传与上下架。"),
        ("FR-03", "AI 商品文案", "根据商品名称、类目、卖点、受众与风格生成标题、卖点和详情文案，并记录生成任务。"),
        ("FR-04", "智能导购", "根据用户需求、预算、候选商品和偏好给出主推荐、理由和备选；支持商品问答与跨商品推荐。"),
        ("FR-05", "评论分析", "对输入评论或指定商品真实评论进行情感占比、关键词、优缺点和改进建议分析。"),
        ("FR-06", "直播脚本", "按商品、平台、时长和语气生成结构化直播/短视频脚本。"),
        ("FR-07", "购物与订单", "购物车增删改选、创建订单、模拟支付、取消、确认收货；商家查看订单并发货。"),
        ("FR-08", "售后与评价", "用户申请退换、提交退货单号、商家完成售后；支持文字、图片和视频评价。"),
        ("FR-09", "个人数据", "维护收货地址、收藏、浏览历史、订单与问答历史。"),
        ("FR-10", "知识库与 RAG", "商家维护规格/FAQ/售后/政策条目，可按商品自动构建；用户问答使用商品上下文。"),
        ("FR-11", "客服会话", "用户发起会话，商家查看线程、回复、标记已读并查看未读计数。"),
        ("FR-12", "采集与运营", "启动京东/苏宁采集任务、查询状态与统计；查看营收看板、用户与问答统计并导出。"),
    ]
    add_table(doc, ["编号", "模块", "需求说明"], functional, [2.0, 3.1, 9.4], 8.8)

    doc.add_heading("4 典型业务流程", level=1)
    doc.add_heading("4.1 用户购买闭环", level=2)
    for step in ["用户登录并浏览/搜索商品。", "将商品加入购物车，选择数量与收货地址。", "创建订单并完成演示支付。", "商家发货，用户确认收货；发生问题时进入售后流程。", "用户提交评价，评论数据可用于商家侧分析。"]:
        add_number(doc, step)
    doc.add_heading("4.2 AI 服务流程", level=2)
    add_figure(doc, ASSETS / "ai-flow.png", "图 2  AI 与 RAG 处理流程")

    doc.add_heading("5 外部接口需求", level=1)
    add_table(doc, ["接口类型", "约定"], [["浏览器界面", "响应式 SPA；桌面端优先，同时适配常见移动宽度"], ["HTTP API", "JSON 为主；统一 /api 前缀；错误返回 error 与 message"], ["认证", "Authorization: Bearer <token>；商家接口额外验证 role=merchant"], ["文件", "头像、商品与评价媒体上传；限制扩展名、大小与存储路径"], ["AI", "OpenAI 兼容 Chat Completions 接口；失败时返回可理解错误，不暴露密钥"], ["数据采集", "外部站点结构变化时允许任务失败并保留状态"]], [3.4, 11.2])
    doc.add_heading("5.1 核心 API 摘要", level=2)
    add_table(doc, ["方法", "路径", "用途", "认证"], [["POST", "/api/auth/login", "登录并签发 JWT", "否"], ["POST", "/api/copy/generate", "生成商品文案", "否"], ["POST", "/api/guide/recommend", "导购推荐", "否"], ["POST", "/api/reviews/analyze", "评论分析", "否"], ["POST", "/api/scripts/live", "直播脚本", "否"], ["GET", "/api/products", "商品列表", "否"], ["POST", "/api/user/orders", "创建订单", "用户"], ["GET", "/api/merchant/dashboard/revenue", "营收看板", "商家"], ["POST", "/api/crawl/start", "启动采集", "商家"], ["POST", "/api/cs/messages", "发起客服消息", "用户"]], [1.6, 5.2, 5.0, 2.6], 8.6)

    doc.add_heading("6 非功能需求", level=1)
    nfrs = [
        ("NFR-01", "安全", "密码使用 Werkzeug scrypt 哈希；生产 JWT_SECRET 至少 32 字符；CORS 仅允许配置来源；不在日志或响应中泄露密钥。"),
        ("NFR-02", "可用性", "登录、加载、提交与失败状态有明确反馈；核心业务在 Mock AI 模式可离线演示。"),
        ("NFR-03", "性能", "前端业务页面异步加载；分页接口避免一次返回全部记录；导出库按需加载。"),
        ("NFR-04", "可靠性", "Redis 不可用时降级为无缓存；PostgreSQL 配置失败不得静默改写到另一个数据库。"),
        ("NFR-05", "可维护性", "后端按蓝图/服务/仓储/模型分层；前端统一 api.ts 与 theme.ts；数据库通过 Alembic 迁移。"),
        ("NFR-06", "兼容性", "支持现代 Chromium/Firefox/Safari；Python 3.11+，Node.js 20.19+ 或 22.12+。"),
        ("NFR-07", "可观测性", "提供 /health、容器日志、采集任务状态和可复现的验证命令。"),
    ]
    add_table(doc, ["编号", "类别", "要求"], nfrs, [2.0, 2.8, 9.8], 8.8)

    doc.add_heading("7 数据、约束与风险", level=1)
    for item in ["默认演示账号只允许本地使用，公网部署必须删除或改密。", "用户上传内容与订单地址属于敏感数据，应限制访问并制定清理策略。", "第三方采集受目标站点结构和规则影响，不能承诺长期稳定。", "AI 输出属于辅助内容，正式发布前需人工审核，避免事实、价格或合规表述错误。", "User.password_plain 字段仅为现有教学实现兼容项，生产化应迁移并删除。"]:
        add_bullet(doc, item)
    doc.add_heading("8 验收标准", level=1)
    add_table(doc, ["验收项", "通过标准"], [["启动", "Docker Compose 可启动 PostgreSQL、Redis、后端与前端；/health 返回正常"], ["认证", "普通用户与商家账号能登录且不能越权进入对方接口"], ["AI", "Mock 模式下四类 AI 工具均可返回结构化结果"], ["交易", "购物车→下单→支付→发货→收货/售后流程可演示"], ["运营", "商品、订单、用户、知识库、客服与看板可访问"], ["质量", "前端生产构建通过；后端自动化测试与静态检查通过"], ["文档", "需求、计划、设计、测试、用户手册与答辩材料齐全且互相一致"]], [3.6, 11.0])
    doc.add_heading("9 需求追踪矩阵", level=1)
    add_table(doc, ["需求", "设计/模块", "验证方式"], [["FR-01", "auth_routes.py / Login.vue", "TC-AUTH-01~04"], ["FR-02", "routes.py / merchant_routes.py / ProductManage.vue", "TC-PROD-01~03"], ["FR-03~06", "AIProvider / OpenAIProvider / 四个 AI 组件", "TC-AI-01~05"], ["FR-07~09", "user_routes.py / shopping.py / 用户端组件", "TC-ORDER-01~05"], ["FR-10", "rag_service.py / knowledge_base.py", "TC-RAG-01~02"], ["FR-11", "customer_service_routes.py", "TC-CS-01~02"], ["FR-12", "crawl_routes.py / RevenueDashboard.vue", "TC-OPS-01~03"]], [2.4, 7.3, 4.9])
    return save_doc(doc, "1.软件需求规约-本项目.docx")


def plan_doc() -> Path:
    title = "项目开发计划"
    doc = new_doc(title, f"{DOC_CODE}-PDP-002")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["项目概述", "范围与交付物", "组织与职责", "生命周期与里程碑", "工作分解", "质量保证", "配置与变更管理", "风险管理", "沟通计划", "验收与收尾"])
    doc.add_heading("1 项目概述", level=1)
    doc.add_paragraph("本项目以课程答辩和可持续演示为目标，建设一个包含真实业务数据模型、双角色界面、AI 工具链和容器部署方案的全栈应用。计划采用短周期迭代：先形成可运行骨架，再补齐业务闭环、质量验证和提交材料。")
    add_callout(doc, "计划基线", "2026-07-07 启动形成骨架；2026-07-15 完成当前实现与提交材料；后续保留答辩演练和缺陷修复缓冲。", BLUE)
    doc.add_heading("1.1 目标", level=2)
    for item in ["完成用户端与商家端关键业务闭环。", "在无真实 AI 密钥时仍可完整演示。", "形成容器化的一键部署和清晰的本地开发路径。", "通过测试、构建、迁移和文档检查形成可验收版本。"]:
        add_bullet(doc, item)
    doc.add_heading("2 范围与交付物", level=1)
    add_table(doc, ["阶段", "主要交付物", "完成定义"], [["策划", "项目计划、分工表、风险清单", "范围、责任、里程碑可追踪"], ["需求", "软件需求规约", "角色、功能、接口、质量属性与验收口径明确"], ["设计", "系统架构与数据库设计", "组件、部署、数据关系与安全决策一致"], ["实现", "Vue 前端、Flask 后端、迁移、容器配置", "核心场景可运行"], ["验证", "测试用例、自动化测试、构建结果", "高优先级用例无阻断缺陷"], ["交付", "用户手册、例会纪要、答辩 PPTX", "材料可打开、内容一致、演示脚本可执行"]], [2.8, 7.0, 4.8])
    doc.add_heading("3 组织与职责", level=1)
    add_table(doc, ["角色", "建议人数", "职责", "成员"], [["项目负责人", "1", "范围、进度、风险、答辩统筹", "待填写"], ["前端开发", "1~2", "交互、页面、API 对接、前端构建", "待填写"], ["后端开发", "1~2", "接口、认证、AI、业务服务", "待填写"], ["数据与测试", "1", "模型、迁移、用例、回归", "待填写"], ["文档与演示", "全员", "文档校核、演示数据、答辩讲解", "待填写"]], [2.8, 2.2, 7.0, 2.6])
    doc.add_heading("4 生命周期与里程碑", level=1)
    milestones = [["M1", "2026-07-07", "项目骨架与运行说明", "已形成"], ["M2", "2026-07-10", "认证、双角色布局、核心模型", "已形成"], ["M3", "2026-07-13", "商品、订单、AI、知识库、客服闭环", "已形成"], ["M4", "2026-07-15", "文档、构建与提交候选版", "当前"], ["M5", "2026-07-18", "答辩演练与缺陷清零", "计划"], ["M6", "2026-07-20", "最终提交与归档", "计划"]]
    add_table(doc, ["里程碑", "日期", "出口工件", "状态"], milestones, [2.0, 3.0, 7.2, 2.4])
    doc.add_heading("5 工作任务分解", level=1)
    add_table(doc, ["WBS", "任务", "责任角色", "依赖", "验收"], [["1.1", "需求梳理与范围冻结", "负责人/全员", "无", "SRS 评审"], ["2.1", "数据库模型与迁移", "后端/数据", "1.1", "迁移可执行"], ["2.2", "API 与认证", "后端", "2.1", "路由与测试"], ["2.3", "AI Provider 与 RAG", "后端", "2.1", "Mock/真实配置可切换"], ["3.1", "用户端页面", "前端", "2.2", "交易主流程"], ["3.2", "商家端页面", "前端", "2.2", "运营主流程"], ["3.3", "前后端联调", "前后端", "2.x/3.x", "无阻断接口错误"], ["4.1", "自动化与手工测试", "测试/全员", "3.3", "测试记录"], ["4.2", "容器部署与验收", "后端/负责人", "4.1", "Compose 健康"], ["5.1", "文档与答辩", "全员", "4.x", "材料齐套"]], [1.5, 4.2, 3.0, 2.3, 3.6], 8.5)
    doc.add_heading("6 质量保证计划", level=1)
    for item in ["每个阶段以可运行工件为出口，不以文档完成代替功能完成。", "后端执行 pytest、ruff、Flask 路由加载与 Alembic current；前端执行类型检查、生产构建和依赖审计。", "关键场景至少覆盖正常流、校验失败、权限拒绝和外部依赖不可用。", "文档中的接口、表名、命令与当前代码交叉核对。", "答辩前按固定演示账号与固定种子数据完成一次全流程彩排。"]:
        add_bullet(doc, item)
    doc.add_heading("7 配置与变更管理", level=1)
    add_table(doc, ["对象", "管理方式"], [["源代码", "Git 分支 + 代码评审；避免提交 .env、数据库、上传文件和构建缓存"], ["数据库", "所有结构变更使用 Alembic；生产启动先 upgrade head"], ["配置", ".env.example 记录变量；密钥只在部署环境注入"], ["需求", "变更需标注影响的接口、模型、页面、测试和文档"], ["发布", "候选版执行验证清单并记录已知问题"]], [3.1, 11.5])
    doc.add_heading("8 风险管理", level=1)
    add_table(doc, ["风险", "概率", "影响", "应对"], [["真实 AI 不可用/配额不足", "中", "高", "默认 Mock；保留 OpenAI 兼容接口切换"], ["采集站点规则变化", "高", "中", "任务状态可见；演示使用种子数据"], ["人员信息未最终确认", "中", "低", "分工表保留待填写字段，提交前统一补齐"], ["数据库配置错误", "中", "高", "区分 SQLite 与 PostgreSQL；禁止静默回退"], ["答辩现场网络不稳定", "中", "高", "提前构建镜像与离线 Mock 演示；准备录屏/截图"], ["文档与代码漂移", "中", "中", "以路由、模型和构建结果为基线复核"]], [4.0, 2.0, 2.0, 6.6])
    doc.add_heading("9 沟通计划", level=1)
    add_table(doc, ["活动", "频率", "参与人", "输出"], [["站会/异步更新", "开发期每日", "项目组", "进度、阻塞、当日目标"], ["项目例会", "每周或里程碑", "全员", "纪要、决策、行动项"], ["需求/设计评审", "阶段出口", "负责人+相关开发", "评审结论与变更"], ["缺陷分诊", "测试期每日", "开发+测试", "优先级与负责人"], ["答辩彩排", "提交前 2 次", "全员", "时间分配与问答清单"]], [3.5, 3.0, 4.2, 4.0])
    doc.add_heading("10 验收与收尾", level=1)
    for step in ["冻结演示数据与候选版本。", "执行后端、前端、迁移、Compose 与手工主流程检查。", "逐份打开交付文件，检查页码、表格、图片与文件名。", "完成答辩彩排，确认各成员讲解模块。", "归档源代码、提交文档、Typst 源文件与最终 PPTX。"]:
        add_number(doc, step)
    return save_doc(doc, "2.项目开发计划-本项目.docx")


def database_doc() -> Path:
    title = "数据库设计说明书"
    doc = new_doc(title, f"{DOC_CODE}-DBD-003")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["引言", "设计目标与约定", "概念结构", "逻辑表设计", "关键字段数据字典", "完整性与索引", "安全与隐私", "迁移与备份"])
    doc.add_heading("1 引言", level=1)
    doc.add_paragraph("本文档描述当前 SQLAlchemy 模型和 Alembic 基线迁移所体现的数据设计。生产环境以 PostgreSQL 16 为目标，本地演示可使用 SQLite；两种数据库共享相同的领域模型。")
    doc.add_heading("2 设计目标与约定", level=1)
    add_table(doc, ["项目", "约定"], [["命名", "表名与字段名使用 snake_case；主键统一 id"], ["时间", "业务时间使用 DateTime；输出 API 时转 ISO 8601"], ["金额", "当前课程实现使用 Float；生产化建议迁移为 Numeric(12,2)"], ["枚举", "订单/售后/角色状态当前为 String，由服务层校验"], ["大字段", "输入输出、快照、关键词与向量数据使用 Text/JSON"], ["迁移", "migrations/versions/0001_schema_baseline.py 为基线；后续变更必须新增 revision"]], [3.4, 11.2])
    doc.add_heading("3 概念结构", level=1)
    add_figure(doc, ASSETS / "data-relationships.png", "图 1  核心实体关系（简化）")
    doc.add_heading("3.1 表清单", level=2)
    tables = [["users", "用户与商家身份"], ["user_addresses", "收货地址"], ["user_favorites", "用户收藏"], ["browse_histories", "浏览历史"], ["products", "商品主数据"], ["reviews", "商品评价"], ["cart_items", "购物车明细"], ["orders", "订单主表与售后状态"], ["order_items", "订单商品快照"], ["generation_tasks", "AI 生成任务"], ["recommendation_logs", "导购推荐记录"], ["knowledge_entries", "商品知识库"], ["qa_records", "用户问答记录"], ["customer_service_messages", "用户—商家客服消息"]]
    add_table(doc, ["表名", "用途"], tables, [6.0, 8.6])
    doc.add_heading("4 逻辑表设计", level=1)
    schemas = {
        "users": [
            ("id", "Integer", "PK, 自增", "内部用户标识"), ("username", "String(100)", "UNIQUE, 非空", "登录账号"), ("password_hash", "String(256)", "非空", "scrypt 哈希"), ("password_plain", "String(100)", "可空", "教学兼容字段，生产化应删除"), ("nickname", "String(100)", "可空", "昵称"), ("role", "String(20)", "默认 user", "user / merchant"), ("avatar/phone/email", "String", "可空", "资料字段"), ("is_active", "Boolean", "默认 true", "是否启用"), ("display_id", "String(64)", "UNIQUE, 可空", "对外展示编号"), ("created_at/updated_at", "DateTime", "自动", "审计时间"),
        ],
        "products": [
            ("id", "Integer", "PK", "商品内部标识"), ("platform/product_id", "String", "可空", "来源平台及平台商品号"), ("name", "String(500)", "非空", "商品名称"), ("category/brand", "String", "可空", "分类与品牌"), ("price/original_price", "Float", "可空", "当前价与原价"), ("selling_points/specs", "Text", "可空", "卖点与规格"), ("image_url/image_urls/videos", "String/JSON", "可空", "媒体资源"), ("detail_url/source_url", "String(500)", "可空", "详情与来源"), ("sales_count/rating/review_count", "Integer/Float", "默认值", "运营统计"), ("display_id", "String(64)", "UNIQUE", "展示编号"), ("is_published", "Boolean", "默认 true", "上架状态"),
        ],
        "orders": [
            ("id/order_no", "Integer/String(64)", "PK/UNIQUE", "内部 ID 与订单号"), ("user_id", "Integer", "非空", "下单用户"), ("status", "String(20)", "默认 pending", "pending/paid/shipped/completed/cancelled/returning/returned"), ("total_amount", "Float", "默认 0", "订单总额"), ("pay_method", "String(20)", "可空", "wechat/alipay"), ("address_snapshot", "Text(JSON)", "可空", "收货地址快照"), ("items_snapshot", "Text(JSON)", "可空", "商品快照"), ("tracking_no", "String(100)", "可空", "发货单号"), ("return_*", "String/DateTime", "可空", "售后原因、状态、单号与时间"), ("created/paid/shipped/completed/cancelled_at", "DateTime", "可空", "生命周期时间"),
        ],
        "knowledge_entries": [
            ("id", "Integer", "PK", "条目 ID"), ("product_id", "Integer", "可空", "关联商品"), ("category", "String(50)", "非空", "spec/faq/after_sale/policy"), ("title/content", "String/Text", "非空", "标题与知识正文"), ("keywords/vector_data", "Text(JSON)", "可空", "关键词与简化向量"), ("is_active", "Boolean", "默认 true", "启用状态"), ("created_at/updated_at", "DateTime", "自动", "审计时间"),
        ],
        "qa_records": [
            ("id", "Integer", "PK", "记录 ID"), ("user_id/product_id", "Integer", "可空", "用户与商品上下文"), ("question/answer", "Text", "非空", "问题与回答"), ("question_type", "String(50)", "默认 auto", "尺码/功能/搭配等"), ("source", "String(50)", "默认 rag", "rag/template/mock"), ("helpful", "Integer", "默认 0", "有用反馈计数"), ("created_at", "DateTime", "自动", "创建时间"),
        ],
        "customer_service_messages": [
            ("id", "Integer", "PK", "消息 ID"), ("user_id", "Integer", "FK users.id", "会话用户"), ("product_id", "Integer", "FK products.id, 可空", "关联商品"), ("sender_role", "String(16)", "默认 user", "user / merchant"), ("content", "Text", "默认空", "消息正文"), ("is_read", "Boolean", "默认 false", "已读标记"), ("created_at", "DateTime", "自动", "发送时间"),
        ],
    }
    for table, fields in schemas.items():
        doc.add_heading(table, level=2)
        add_table(doc, ["字段", "类型", "约束", "说明"], fields, [4.0, 3.0, 3.6, 4.0], 8.5)
    doc.add_heading("5 关系与完整性", level=1)
    add_table(doc, ["父实体", "子实体", "关系", "策略"], [["products", "reviews", "1:N", "删除商品时 ORM 级联删除评论"], ["products", "generation_tasks", "1:N", "任务保留关联"], ["orders", "order_items", "1:N", "删除订单时级联明细"], ["users", "customer_service_messages", "1:N", "数据库外键"], ["products", "customer_service_messages", "1:N", "商品可空外键"], ["users/products", "购物车、收藏、历史、问答", "1:N", "部分关系由业务层保证"]], [3.2, 4.6, 2.2, 4.6])
    doc.add_heading("5.1 索引建议", level=2)
    for item in ["users.username、users.display_id、products.display_id、orders.order_no 保持唯一索引。", "customer_service_messages.user_id 已建立索引；建议追加 (user_id, created_at)。", "建议为 products(category, is_published)、orders(user_id, status, created_at)、reviews(product_id, created_at) 建立组合索引。", "知识库检索量增长后，vector_data 应迁移到专用向量类型/索引或外部向量库。"]:
        add_bullet(doc, item)
    doc.add_heading("6 安全与隐私", level=1)
    for item in ["只保存密码哈希；现有 password_plain 字段不得用于生产，迁移前先清空历史值。", "订单地址快照、手机号、邮箱与客服内容按个人信息处理，限制商家访问范围并配置备份保留期。", "测试与日志使用脱敏数据，不输出 JWT、AI_API_KEY、数据库密码。", "上传文件使用随机文件名、扩展名白名单和体积限制，禁止任意路径访问。"]:
        add_bullet(doc, item)
    doc.add_heading("7 迁移、备份与恢复", level=1)
    add_table(doc, ["场景", "操作"], [["初始化", "alembic upgrade head；本地 SQLite 可由 init_db() 保证首次可用"], ["结构变更", "alembic revision --autogenerate -m <说明>，人工复核后 upgrade"], ["检查", "alembic current；确认数据库 URL 与预期环境一致"], ["备份", "PostgreSQL 使用 pg_dump；文件放 data/backups/，不提交 Git"], ["恢复", "新库先恢复备份，再执行剩余迁移；恢复后做账号、商品、订单抽查"], ["回滚", "优先向前修复；破坏性降级前必须备份并演练"]], [3.0, 11.6])
    return save_doc(doc, "3.数据库设计说明书-本项目.docx")


def architecture_doc() -> Path:
    title = "系统构架设计说明书"
    doc = new_doc(title, f"{DOC_CODE}-SAD-004")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["简介", "架构目标与约束", "系统上下文", "逻辑视图", "运行时视图", "部署视图", "接口与数据视图", "安全设计", "质量属性", "关键决策与演进"])
    doc.add_heading("1 简介", level=1)
    doc.add_paragraph("本文档使用上下文、逻辑、运行时、部署和数据视图描述系统，并记录对演示可用性、部署方式、AI 可替换性和数据一致性影响最大的架构决策。")
    doc.add_heading("2 架构目标与约束", level=1)
    add_table(doc, ["目标/约束", "设计响应"], [["离线可演示", "Mock AI + SQLite；Redis 失败自动降级"], ["生产可部署", "Docker Compose 组织 PostgreSQL、Redis、Gunicorn 和 Nginx"], ["双角色隔离", "JWT 身份 + role=merchant 后端校验；前端按角色加载布局"], ["维护成本可控", "后端蓝图/服务/模型分层，前端统一 API 与主题入口"], ["第三方不稳定", "AI Provider 抽象；采集任务化并暴露状态"], ["课程时间有限", "单体后端与单 SPA，避免过早拆分微服务"]], [4.0, 10.6])
    doc.add_heading("3 系统上下文", level=1)
    add_figure(doc, ASSETS / "system-architecture.png", "图 1  系统上下文与部署边界")
    doc.add_heading("4 逻辑视图", level=1)
    add_table(doc, ["层/组件", "职责", "主要实现"], [["表现层", "登录、用户前台、商家后台、AI 工具与反馈", "Vue 3、Ant Design Vue、Tailwind、Motion"], ["API 层", "请求路由、校验、认证、错误响应", "Flask 蓝图、Pydantic、JWT、CORS"], ["业务服务层", "AI、RAG、认证、上传、采集编排", "backend/services、backend/crawler"], ["数据访问层", "会话、仓储、ORM 模型、迁移", "SQLAlchemy、Alembic"], ["基础设施层", "数据库、缓存、外部 AI、Nginx", "PostgreSQL/SQLite、Redis、OpenAI、Docker"]], [3.0, 6.0, 5.6])
    doc.add_heading("4.1 后端蓝图", level=2)
    add_table(doc, ["蓝图", "职责"], [["api_bp", "AI 能力、公共商品、评论与统计"], ["auth_bp", "登录、注册、密码、个人资料、头像"], ["merchant_bp", "商品、知识库、订单、用户、运营看板"], ["user_bp", "购物车、订单、地址、收藏、历史、问答、评价"], ["crawl_bp", "采集预设、启动、状态、任务与统计"], ["cs_bp", "用户消息、商家线程、回复、已读与未读计数"], ["docs_bp", "API 文档页面与 OpenAPI JSON"]], [3.4, 11.2])
    doc.add_heading("4.2 前端组件", level=2)
    add_table(doc, ["区域", "主要组件"], [["入口", "App.vue、Login.vue、api.ts、theme.ts"], ["商家端", "MerchantLayout、RevenueDashboard、ProductManage、OrderManage、UserManage、KnowledgeBaseManage、QAStats、MerchantCustomerService"], ["用户端", "UserLayout、ProductBrowse、ShoppingCart、MyOrders、UserProfile、CustomerService"], ["AI 工具", "CopyGenerator、ReviewAnalyzer、LiveScriptGenerator、ChatWidget"]], [3.4, 11.2])
    doc.add_heading("5 运行时视图", level=1)
    doc.add_heading("5.1 登录与授权", level=2)
    for step in ["前端提交用户名、密码和选择的角色。", "认证服务校验密码哈希，并拒绝角色与账号不匹配。", "后端签发 24 小时 JWT，前端保存会话并调用 /api/auth/me。", "受保护接口校验 JWT；商家接口进一步验证角色。", "401/403 时前端清理无效会话并提示重新登录。"]:
        add_number(doc, step)
    doc.add_heading("5.2 AI / RAG", level=2)
    add_figure(doc, ASSETS / "ai-flow.png", "图 2  AI 请求的运行时流程")
    doc.add_heading("5.3 订单状态", level=2)
    add_table(doc, ["当前状态", "允许动作", "目标状态"], [["pending", "用户支付 / 用户或商家取消", "paid / cancelled"], ["paid", "商家发货", "shipped"], ["shipped", "用户确认收货", "completed"], ["paid/shipped/completed", "用户申请退换", "returning"], ["returning", "用户提交退货单号 / 商家完成售后", "returned / refunded/exchanged"]], [3.5, 6.8, 4.3])
    doc.add_heading("6 部署视图", level=1)
    add_table(doc, ["容器", "端口/网络", "持久化", "健康/启动"], [["frontend", "宿主 8080 → Nginx 80", "静态构建", "依赖 backend；代理 /api 与 /health"], ["backend", "宿主 127.0.0.1:8000", "uploads 可挂载", "先 Alembic，再 Gunicorn"], ["postgres", "容器 5432", "postgres_data volume", "数据库健康检查"], ["redis", "容器 6379", "可选持久化", "连接失败业务降级"]], [3.0, 4.0, 3.8, 3.8])
    doc.add_heading("7 接口与数据视图", level=1)
    doc.add_paragraph("接口采用 JSON/REST 风格，以资源和业务动作组合命名。OpenAPI 页面覆盖核心公共 AI 接口，完整契约仍以 Flask 路由和请求模型为准。数据关系详见《数据库设计说明书》。")
    add_figure(doc, ASSETS / "data-relationships.png", "图 3  数据关系概览")
    doc.add_heading("8 安全设计", level=1)
    add_table(doc, ["威胁", "控制"], [["弱口令/口令泄露", "scrypt 哈希；生产修改默认账号与密码"], ["令牌伪造", "至少 32 字符随机 JWT_SECRET；HTTPS；短期令牌"], ["越权", "服务端 JWT 与角色校验，不依赖前端菜单隐藏"], ["跨域滥用", "CORS_ORIGINS 白名单"], ["恶意上传", "扩展名、大小、随机文件名与受控目录"], ["密钥泄露", ".env 不入库，日志不打印 AI_API_KEY"], ["敏感数据扩散", "限制商家查询范围、备份加密与定期清理"]], [4.0, 10.6])
    doc.add_heading("9 质量属性场景", level=1)
    add_table(doc, ["属性", "场景", "响应"], [["可用性", "Redis 宕机", "请求继续，无缓存；记录警告"], ["可部署性", "新环境上线", "Compose 构建、迁移、健康检查后可访问"], ["性能", "用户仅访问登录页", "业务页面异步加载，避免下载全后台"], ["可修改性", "替换 AI 服务", "实现 AIProvider，不改路由与前端契约"], ["可靠性", "PostgreSQL 配置错误", "明确失败，不静默写入 SQLite"], ["可测试性", "无外网/无密钥", "Mock AI 返回确定性结果"]], [3.0, 5.4, 6.2])
    doc.add_heading("10 关键决策与演进", level=1)
    add_table(doc, ["决策", "理由", "未来触发条件"], [["模块化单体", "课程规模下交付与调试成本最低", "团队/流量显著增长再拆服务"], ["OpenAI 兼容抽象", "降低供应商绑定并保留本地 Mock", "引入多模型路由与用量计费"], ["数据库双模式", "SQLite 便于演示，PostgreSQL 适合部署", "生产统一 PostgreSQL 后删除兼容分支"], ["Redis 可选", "缓存不应阻断核心业务", "引入分布式任务后提升为必选依赖"], ["Nginx 同源代理", "简化 CORS 与前端配置", "多域部署时显式配置 API 基址"]], [4.0, 6.0, 4.6])
    return save_doc(doc, "3.系统构架设计说明书-本项目.docx")


def test_doc() -> Path:
    title = "测试用例"
    doc = new_doc(title, f"{DOC_CODE}-TST-005")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["测试目标", "测试环境", "准入与准出", "用例总览", "详细用例", "自动化覆盖", "缺陷分级", "回归清单"])
    doc.add_heading("1 测试目标", level=1)
    doc.add_paragraph("验证认证隔离、AI 工具、商品与交易闭环、知识库、客服、采集任务、部署降级和前端构建，重点检查数据一致性、权限边界与失败反馈。")
    doc.add_heading("2 测试环境", level=1)
    add_table(doc, ["项目", "环境"], [["浏览器", "最新版 Chromium / Firefox，桌面宽度 1440px，移动宽度 390px"], ["前端", "Node.js 20.19+ 或 22.12+；Vite 开发与生产构建"], ["后端", "Python 3.11+；Flask 3；pytest"], ["数据", "SQLite 演示环境；PostgreSQL 16 集成环境；Redis 7 可用与不可用两种场景"], ["AI", "Mock AI 必测；OpenAI 兼容接口作为可选集成测试"], ["账号", "merchant / merchant123；user / user123（仅本地）"]], [3.2, 11.4])
    doc.add_heading("3 准入与准出", level=1)
    add_table(doc, ["类型", "条件"], [["准入", "依赖安装完成；迁移可执行；种子数据可创建；需求与接口处于基线状态"], ["准出", "P0/P1 缺陷为 0；关键主流程通过；自动化测试、前端构建、迁移与 Compose 检查通过；遗留问题有说明"]], [3.2, 11.4])
    cases = [
        ("TC-AUTH-01", "用户登录", "输入 user/user123 并选择用户", "签发 JWT，进入用户端", "P0"),
        ("TC-AUTH-02", "角色不匹配", "merchant 账号选择用户身份", "拒绝并提示切换身份", "P0"),
        ("TC-AUTH-03", "错误密码", "输入错误密码", "401；不返回敏感信息", "P1"),
        ("TC-AUTH-04", "商家越权", "用户 JWT 调用 /api/merchant/products", "403", "P0"),
        ("TC-AI-01", "文案生成", "提交完整商品与风格参数", "返回标题/卖点/详情且任务可记录", "P0"),
        ("TC-AI-02", "文案校验", "缺少必填商品名", "400，字段错误清晰", "P1"),
        ("TC-AI-03", "导购推荐", "需求+预算+候选商品", "主推荐、理由与备选结构完整", "P0"),
        ("TC-AI-04", "评论分析", "提交多条正负评论", "情感、关键词、优缺点与建议", "P1"),
        ("TC-AI-05", "直播脚本", "选择商品、平台和时长", "脚本按段落/节奏输出", "P1"),
        ("TC-PROD-01", "商品分页筛选", "关键词+分类+分页", "列表、总数和页码一致", "P0"),
        ("TC-PROD-02", "商家新增商品", "提交合法商品信息", "创建成功并可查询", "P0"),
        ("TC-PROD-03", "不存在商品", "访问无效 ID", "404", "P1"),
        ("TC-ORDER-01", "购物车", "新增、改数量、取消选择、删除", "金额与项目状态正确", "P0"),
        ("TC-ORDER-02", "下单支付", "从购物车选择商品和地址", "订单快照正确，pending→paid", "P0"),
        ("TC-ORDER-03", "发货收货", "商家发货，用户确认", "paid→shipped→completed", "P0"),
        ("TC-ORDER-04", "售后", "申请退货、提交单号、商家完成", "状态与时间字段一致", "P0"),
        ("TC-RAG-01", "知识库问答", "建立商品知识后提问", "回答引用相关上下文并记录", "P1"),
        ("TC-CS-01", "客服会话", "用户发消息，商家回复并标记已读", "线程顺序、角色与未读数正确", "P1"),
        ("TC-OPS-01", "采集任务", "启动预设关键词采集并查询状态", "任务 ID、进度和失败信息可见", "P2"),
        ("TC-NFR-01", "缓存降级", "停止 Redis 后请求 AI/商品接口", "核心请求继续，日志记录降级", "P1"),
    ]
    doc.add_heading("4 用例总览", level=1)
    add_table(doc, ["编号", "场景", "关键输入/步骤", "预期结果", "优先级"], cases, [2.3, 3.0, 4.5, 5.0, 1.5], 7.9)
    doc.add_heading("5 代表性详细用例", level=1)
    details = [
        ("TC-AUTH-02 角色不匹配", ["前置：默认商家账号存在。", "步骤：打开登录页；选择‘用户’；输入 merchant/merchant123；提交。", "预期：HTTP 401；页面明确提示账号属于商家管理员；不保存 token；不进入任何业务布局。"]),
        ("TC-AI-01 文案生成", ["前置：AI_PROVIDER=mock；后端可用。", "步骤：输入商品名称、类目、卖点、目标受众与风格；点击生成。", "预期：显示加载态；返回标题、核心卖点和详情文案；页面可复制；生成任务状态为 completed；无密钥也可完成。"]),
        ("TC-ORDER-02 下单支付", ["前置：用户登录；购物车至少一件选中商品；存在默认地址。", "步骤：进入购物车结算；核对商品/数量/金额；创建订单；选择演示支付方式并支付。", "预期：订单号唯一；地址与商品快照完整；总金额等于明细价格×数量之和；状态 pending→paid；购物车对应项目清理。"]),
        ("TC-NFR-01 缓存降级", ["前置：配置 REDIS_URL，但 Redis 服务停止。", "步骤：启动后端；请求核心 AI 与商品接口。", "预期：服务不因连接异常退出；请求继续执行；日志仅记录警告；恢复 Redis 后新请求可重新使用缓存。"]),
    ]
    for heading, items in details:
        doc.add_heading(heading, level=2)
        for item in items:
            add_bullet(doc, item)
    doc.add_heading("6 自动化覆盖与当前验证", level=1)
    add_table(doc, ["检查", "位置/命令", "当前说明"], [["健康、能力、文案与校验", "tests/test_api.py", "已有 pytest 用例"], ["JWT、密码哈希、CORS", "tests/test_auth.py", "已有 pytest 用例"], ["商品数据库用例", "tests/test_api.py", "需 PostgreSQL，当前标记跳过"], ["前端类型与构建", "npm run build", "2026-07-15 在当前工作区通过"], ["后端全量回归", "python -m pytest -q", "交付前在项目 Python 环境执行"], ["静态检查", "ruff check backend tests migrations scripts", "交付前执行"], ["迁移", "alembic upgrade head && alembic current", "SQLite/PostgreSQL 各执行一次"], ["部署", "docker compose config --quiet", "交付前执行"]], [3.4, 6.2, 5.0])
    doc.add_heading("7 缺陷分级", level=1)
    add_table(doc, ["等级", "定义", "处理"], [["P0 阻断", "无法启动、数据破坏、认证绕过、主流程完全不可用", "立即修复，禁止提交"], ["P1 严重", "核心功能错误、状态不一致、明显敏感信息泄露", "候选版前修复"], ["P2 一般", "次要功能或兼容问题，有可接受绕行", "评估后修复或记录"], ["P3 建议", "文案、细节与体验优化", "进入后续迭代"]], [2.5, 8.0, 4.1])
    doc.add_heading("8 回归清单", level=1)
    for item in ["两种角色登录与越权访问", "四类 AI 工具在 Mock 模式输出", "商品筛选与商家 CRUD", "购物车、下单、支付、发货、收货与售后", "知识库问答与客服线程", "Redis 降级与数据库配置失败行为", "前端构建、后端测试、迁移、Compose 配置", "默认账号、环境变量和上传目录安全"]:
        add_bullet(doc, item)
    return save_doc(doc, "4.测试用例-本项目.docx")


def meeting_doc() -> Path:
    title = "项目例会纪要"
    doc = new_doc(title, f"{DOC_CODE}-MIN-006", subtitle="里程碑评审与提交材料准备")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["会议基本信息", "进展回顾", "议题与决策", "行动项", "风险与问题", "会议签到"])
    doc.add_heading("1 会议基本信息", level=1)
    add_table(doc, ["日期", "时间", "地点/方式"], [[TODAY.isoformat(), "19:30—20:15", "项目组会议 / 线上或教室（待填写）"]], [4.0, 4.0, 6.6])
    add_table(doc, ["主题", "主持人", "记录人", "出席者"], [["提交候选版评审与答辩材料准备", "待填写", "待填写", "项目组全体成员（姓名待补充）"]], [5.6, 3.0, 3.0, 3.0])
    doc.add_heading("2 进展回顾", level=1)
    add_table(doc, ["工作项", "完成情况", "结论"], [["前后端业务骨架", "Vue 3 + Flask 主流程已形成", "进入验收"], ["双角色与认证", "用户/商家布局、JWT、资料与权限接口已实现", "补充越权回归"], ["AI 工具", "文案、导购、评论、脚本支持 Mock/OpenAI", "答辩使用 Mock 保底"], ["交易与运营", "商品、订单、售后、知识库、客服、采集与看板已形成", "冻结演示数据"], ["工程化", "Compose、Nginx、Gunicorn、Alembic 与运行说明已具备", "补做全链路检查"], ["提交材料", "按参考模板创建需求、计划、设计、测试、手册与 PPT", "统一编号与术语"]], [4.0, 7.0, 3.6])
    doc.add_heading("3 议题与决策", level=1)
    decisions = [["D-01", "演示依赖", "现场默认使用 Mock AI 与本地种子数据；真实 AI 作为可选展示。", "通过"], ["D-02", "部署方式", "优先 Docker Compose 一键启动；保留前后端双终端开发方式。", "通过"], ["D-03", "人员信息", "不在材料中虚构成员姓名；分工和签字栏由项目组提交前补齐。", "通过"], ["D-04", "答辩结构", "按背景→方案→功能→技术→难点→测试→分工→总结讲解。", "通过"], ["D-05", "质量门槛", "P0/P1 缺陷清零，前端构建、后端测试、迁移和 Compose 检查通过后提交。", "通过"]]
    add_table(doc, ["编号", "议题", "决策", "状态"], decisions, [2.0, 3.4, 7.2, 2.0])
    doc.add_heading("4 行动项", level=1)
    actions = [["A-01", "补齐成员姓名、学号、联系方式与分工比例", "项目负责人", "2026-07-16", "待完成"], ["A-02", "执行后端测试、静态检查、迁移与 Compose 验证", "后端/测试", "2026-07-17", "待完成"], ["A-03", "冻结演示数据并彩排用户购买与商家发货流程", "前端/测试", "2026-07-17", "待完成"], ["A-04", "逐份校对提交文档与页码、文件名", "文档负责人", "2026-07-18", "待完成"], ["A-05", "完成两轮答辩彩排与问答准备", "全员", "2026-07-19", "待完成"]]
    add_table(doc, ["编号", "行动", "负责人", "截止日期", "状态"], actions, [1.8, 6.2, 3.0, 2.8, 2.0], 8.5)
    doc.add_heading("5 风险与问题", level=1)
    add_table(doc, ["问题/风险", "影响", "处置"], [["人员资料尚未提供", "提交表与封面信息不完整", "使用待填写字段，提交前由负责人一次性补齐"], ["现场网络不可控", "真实 AI 与采集可能失败", "Mock AI、种子数据、离线构建与备份截图"], ["后端全量测试未在本次文档生成环境执行", "不能把文档生成当作功能验收", "在项目依赖环境按测试清单执行并记录结果"]], [4.5, 4.2, 6.3])
    doc.add_heading("6 会议签到", level=1)
    rows = [[i, "", role, "", ""] for i, role in enumerate(["项目负责人", "前端开发", "后端开发", "数据/测试", "文档/演示", "其他"], 1)]
    add_table(doc, ["序号", "姓名", "角色", "签名", "备注"], rows, [1.5, 3.2, 3.6, 3.2, 3.1])
    return save_doc(doc, "5.项目例会纪要-本项目.docx")


def manual_doc() -> Path:
    title = "用户使用手册"
    doc = new_doc(title, f"{DOC_CODE}-UM-007", subtitle="用户端、商家端与系统维护者操作指南")
    add_front_matter(doc, f"{PROJECT}{title}", contents=["阅读说明", "快速启动", "登录与账号", "普通用户操作", "商家端操作", "AI 工具", "数据采集", "导出", "常见问题", "安全与维护"])
    doc.add_heading("1 阅读说明", level=1)
    doc.add_paragraph("本手册面向普通用户、商家管理员和演示维护者。界面名称以当前 Vue 应用为准；API 细节可访问 /api/docs。默认账号仅供本地课程演示。")
    add_callout(doc, "最快演示", "执行 docker compose up --build -d，然后访问 http://localhost:8080。无需 AI 密钥即可使用 Mock AI。", MINT)
    doc.add_heading("2 快速启动", level=1)
    doc.add_heading("2.1 Docker Compose（推荐）", level=2)
    for step in ["确认 Docker Engine/Desktop 与 Docker Compose 可用。", "可选：复制 .env.example 为 .env，并修改密码、JWT_SECRET 或 AI 配置。", "在项目根目录运行 docker compose up --build -d。", "运行 docker compose ps，确认服务正常。", "访问应用 http://localhost:8080；健康检查 http://localhost:8000/health；API 文档 http://localhost:8000/api/docs。"]:
        add_number(doc, step)
    doc.add_heading("2.2 本地开发", level=2)
    add_table(doc, ["终端", "命令"], [["后端", "python -m venv .venv；激活后 pip install -r requirements-dev.txt；alembic upgrade head；python -m flask --app backend.app run --host=0.0.0.0 --port=8000"], ["前端", "cd frontend；npm ci；npm run dev；访问 http://localhost:5173"]], [3.0, 11.6])
    doc.add_heading("3 登录与账号", level=1)
    add_table(doc, ["角色", "用户名", "密码", "用途"], [["商家", "merchant", "merchant123", "进入运营后台"], ["用户", "user", "user123", "进入购物前台"]], [3.0, 3.6, 3.6, 4.4])
    for step in ["打开应用，选择‘普通用户’或‘商家管理员’身份。", "输入用户名和密码。系统会校验账号真实角色；选择错误会提示切换。", "登录后可在个人资料中修改昵称、手机号、邮箱和头像。", "离开公共设备前使用退出登录。"]:
        add_number(doc, step)
    doc.add_heading("4 普通用户操作", level=1)
    sections = [
        ("4.1 浏览商品", ["在商品页按分类或关键词筛选。", "打开商品详情查看价格、规格、评分、媒体与评论。", "收藏商品或写入浏览历史，便于在个人中心再次访问。"]),
        ("4.2 购物车与结算", ["点击加入购物车；在购物车调整数量和选中状态。", "确认至少存在一个收货地址；可设置默认地址。", "选择商品结算，核对金额和备注后创建订单。"]),
        ("4.3 订单", ["待支付：选择演示支付方式完成支付，或取消。", "已发货：查看物流单号并确认收货。", "订单详情保留商品与地址快照，后续商品修改不影响历史订单。"]),
        ("4.4 售后与评价", ["在符合条件的订单中申请退换并填写原因。", "按页面提示提交退货单号，等待商家处理。", "订单完成后提交星级、文字以及可选图片/视频评价。"]),
        ("4.5 智能问答与客服", ["在商品相关页面向 AI 提问尺码、功能、搭配等问题。", "系统优先使用商品与知识库上下文并保存问答历史。", "需要人工协助时进入客服会话，商家回复会显示在同一线程。"]),
    ]
    for heading, items in sections:
        doc.add_heading(heading, level=2)
        for item in items:
            add_bullet(doc, item)
    doc.add_heading("5 商家端操作", level=1)
    merchant = [
        ("5.1 营收看板", ["查看收入、订单、用户和趋势数据。", "按页面提供的导出按钮保存 JPG 或 XLSX；导出模块按需加载。"]),
        ("5.2 商品管理", ["新增或编辑商品名称、分类、品牌、价格、卖点、规格和媒体。", "使用上传入口批量导入数据；操作前检查字段模板。", "删除商品前确认订单/评论等历史依赖；优先使用上下架。"]),
        ("5.3 订单与售后", ["按状态筛选订单并查看详情。", "对已支付订单填写快递单号并发货。", "对退换申请核对退货单号并完成退款/换货处理。"]),
        ("5.4 用户管理", ["查看、创建、编辑或停用用户。", "仅在授权范围内处理手机号、邮箱等个人信息。"]),
        ("5.5 知识库与问答统计", ["新增 spec、faq、after_sale、policy 等知识条目。", "可按商品自动构建基础知识，再人工校对。", "查看问答记录、来源与统计，发现高频问题后补充知识。"]),
        ("5.6 客服", ["在线程列表查看各用户最近消息与未读数。", "进入线程查看完整上下文并回复；阅读后标记已读。"]),
    ]
    for heading, items in merchant:
        doc.add_heading(heading, level=2)
        for item in items:
            add_bullet(doc, item)
    doc.add_heading("6 AI 工具", level=1)
    add_table(doc, ["工具", "输入", "输出/建议"], [["商品文案", "商品、类目、卖点、受众、风格", "标题、卖点与详情；发布前人工复核"], ["智能导购", "需求、预算、候选商品、偏好", "主推荐、理由、备选与购买建议"], ["评论分析", "评论文本或商品 ID", "情感占比、关键词、优缺点、改进建议"], ["直播脚本", "商品、平台、时长、语气", "开场、卖点、互动、转化与收尾脚本"]], [3.2, 5.7, 5.7])
    add_callout(doc, "AI 输出责任", "生成内容可能出现事实偏差、夸大宣传或过时信息；价格、功效、售后和合规表述必须由运营人员审核后发布。", BLUE)
    doc.add_heading("7 数据采集", level=1)
    for step in ["以商家身份进入采集功能，先查看预设关键词。", "选择平台/关键词并启动任务，记录返回的 task_id。", "查看任务状态或任务列表；若目标站点结构变化，任务可能失败。", "核对采集商品与评论，修正分类、价格和媒体后再上架。"]:
        add_number(doc, step)
    doc.add_heading("8 常见问题", level=1)
    add_table(doc, ["现象", "处理"], [["页面能开但 API 报错", "检查后端 8000 端口与 /health；开发环境确认 Vite 代理"], ["无法登录", "确认选择的身份与账号角色一致；仅本地使用默认账号"], ["AI 提示缺少密钥", "将 AI_PROVIDER 设为 mock；或配置 AI_API_KEY/AI_MODEL/AI_BASE_URL"], ["Redis 连接失败", "核心业务会降级；需要缓存时启动 redis 容器并检查 REDIS_URL"], ["数据库连接失败", "检查 DATABASE_URL；PostgreSQL 失败不会自动写到 SQLite"], ["采集没有结果", "检查任务状态、网络、站点规则和日志；演示使用种子数据"], ["前端改动未生效", "开发模式重启 Vite；容器模式重新 docker compose up --build -d"]], [4.3, 10.3])
    doc.add_heading("9 安全与维护", level=1)
    for item in ["公网部署前替换 JWT_SECRET、数据库密码和默认账号。", "限制 CORS_ORIGINS，使用 HTTPS，并将后端端口仅绑定本机或内网。", "定期执行数据库备份；不要把 .env、数据库文件、用户上传和备份提交 Git。", "升级模型结构时使用 Alembic，不直接修改生产表。", "更新依赖后运行 pytest、ruff、npm run build、npm audit 与 Compose 配置检查。"]:
        add_bullet(doc, item)
    doc.add_heading("10 维护命令速查", level=1)
    add_table(doc, ["目的", "命令"], [["查看日志", "docker compose logs -f backend"], ["重建", "docker compose up --build -d"], ["停止保留数据", "docker compose down"], ["迁移", "alembic upgrade head && alembic current"], ["后端测试", "python -m pytest -q"], ["静态检查", "ruff check backend tests migrations scripts"], ["前端构建", "cd frontend && npm run build"], ["健康检查", "curl http://localhost:8000/health"]], [4.0, 10.6])
    return save_doc(doc, "6.用户使用手册-本项目.docx")


def assignments_xlsx() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "项目组成员分工"
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A6"
    ws.merge_cells("A1:H1")
    ws["A1"] = PROJECT
    ws["A2"] = "项目组成员分工表"
    ws.merge_cells("A2:H2")
    ws["A3"] = "说明：姓名、学号与实际工作量由项目组提交前补齐；完成度与工时为可编辑项。"
    ws.merge_cells("A3:H3")
    ws["A1"].font = Font(name="Microsoft YaHei", size=18, bold=True, color=WHITE)
    ws["A1"].fill = PatternFill("solid", fgColor=NAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws["A2"].font = Font(name="Microsoft YaHei", size=13, bold=True, color=BLUE)
    ws["A2"].alignment = Alignment(horizontal="center")
    ws["A3"].font = Font(name="Microsoft YaHei", size=10, color=MUTED)
    ws["A3"].alignment = Alignment(horizontal="left")
    ws.row_dimensions[1].height = 34
    ws.row_dimensions[2].height = 25
    ws.row_dimensions[3].height = 23

    headers = ["序号", "姓名", "学号", "项目角色", "负责模块", "主要工作成果", "完成度", "备注"]
    for col, value in enumerate(headers, 1):
        c = ws.cell(5, col, value)
        c.font = Font(name="Microsoft YaHei", size=10, bold=True, color=WHITE)
        c.fill = PatternFill("solid", fgColor=NAVY)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    rows = [
        [1, "待填写", "待填写", "项目负责人", "范围、计划、集成、答辩", "里程碑管理、风险与交付统筹", 0, ""],
        [2, "待填写", "待填写", "前端开发", "用户端与通用交互", "登录、商品浏览、购物车、订单、个人中心", 0, ""],
        [3, "待填写", "待填写", "前端开发", "商家端与 AI 工具", "看板、商品/订单/知识库/客服、AI 页面", 0, ""],
        [4, "待填写", "待填写", "后端开发", "API、认证与业务服务", "Flask 蓝图、JWT、订单、客服、上传", 0, ""],
        [5, "待填写", "待填写", "数据与 AI", "数据库、RAG、采集", "SQLAlchemy/Alembic、AI Provider、知识库、爬虫", 0, ""],
        [6, "待填写", "待填写", "测试与文档", "测试、部署与提交材料", "pytest/构建/Compose 验证、文档和答辩材料", 0, ""],
    ]
    for r, row in enumerate(rows, 6):
        for col, value in enumerate(row, 1):
            c = ws.cell(r, col, value)
            c.font = Font(name="Microsoft YaHei", size=10, color=INK)
            c.alignment = Alignment(horizontal="center" if col in (1, 2, 3, 4, 7) else "left", vertical="center", wrap_text=True)
            if r % 2 == 1:
                c.fill = PatternFill("solid", fgColor=PALE)
        ws.cell(r, 7).number_format = "0%"
        ws.row_dimensions[r].height = 46

    thin = Side(style="thin", color=LINE)
    for row in ws.iter_rows(min_row=5, max_row=11, min_col=1, max_col=8):
        for c in row:
            c.border = Border(bottom=thin)
    widths = {"A": 8, "B": 13, "C": 16, "D": 16, "E": 28, "F": 38, "G": 12, "H": 18}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    dv_role = DataValidation(type="list", formula1='"项目负责人,前端开发,后端开发,数据与 AI,测试与文档,其他"', allow_blank=True)
    ws.add_data_validation(dv_role)
    dv_role.add("D6:D30")
    dv_progress = DataValidation(type="decimal", operator="between", formula1="0", formula2="1", allow_blank=True)
    ws.add_data_validation(dv_progress)
    dv_progress.add("G6:G30")
    ws.conditional_formatting.add("G6:G30", FormulaRule(formula=["G6=1"], fill=PatternFill("solid", fgColor="DCFCE7")))
    ws.auto_filter.ref = "A5:H11"
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.page_setup.orientation = "landscape"
    ws.print_area = "A1:H11"

    summary = wb.create_sheet("模块责任矩阵")
    summary.sheet_view.showGridLines = False
    summary.merge_cells("A1:G1")
    summary["A1"] = "模块责任矩阵（R=主责，C=协作，V=验证）"
    summary["A1"].font = Font(name="Microsoft YaHei", size=15, bold=True, color=WHITE)
    summary["A1"].fill = PatternFill("solid", fgColor=NAVY)
    summary["A1"].alignment = Alignment(horizontal="center")
    matrix_headers = ["模块", "负责人", "前端1", "前端2", "后端", "数据/AI", "测试/文档"]
    matrix = [
        ["需求与计划", "R", "C", "C", "C", "C", "V"],
        ["用户端", "C", "R", "C", "C", "C", "V"],
        ["商家端", "C", "C", "R", "C", "C", "V"],
        ["认证与 API", "C", "C", "C", "R", "C", "V"],
        ["数据库与迁移", "C", "", "", "C", "R", "V"],
        ["AI 与 RAG", "C", "C", "C", "R", "R", "V"],
        ["测试与部署", "C", "C", "C", "C", "C", "R"],
        ["文档与答辩", "R", "C", "C", "C", "C", "R"],
    ]
    for col, h in enumerate(matrix_headers, 1):
        c = summary.cell(3, col, h)
        c.font = Font(name="Microsoft YaHei", size=10, bold=True, color=WHITE)
        c.fill = PatternFill("solid", fgColor=BLUE)
        c.alignment = Alignment(horizontal="center")
    for r, row in enumerate(matrix, 4):
        for col, value in enumerate(row, 1):
            c = summary.cell(r, col, value)
            c.font = Font(name="Microsoft YaHei", size=10, bold=col > 1 and value == "R", color=INK)
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = Border(bottom=thin)
            if value == "R":
                c.fill = PatternFill("solid", fgColor="DBEAFE")
            elif value == "V":
                c.fill = PatternFill("solid", fgColor="ECFDF5")
    for col in "ABCDEFG":
        summary.column_dimensions[col].width = 18 if col != "A" else 24
    summary.freeze_panes = "B4"
    summary.sheet_properties.pageSetUpPr.fitToPage = True
    summary.page_setup.fitToWidth = 1
    summary.page_setup.orientation = "landscape"
    path = OUT / "项目组成员分工-本项目.xlsx"
    wb.save(path)
    return path


def build_manifest(paths: list[Path]) -> None:
    manifest = {
        "project": PROJECT,
        "generated_at": TODAY.isoformat(),
        "files": [{"name": p.name, "bytes": p.stat().st_size} for p in paths],
        "notes": ["人员姓名、学号、签字为待填写项。", "内容按当前仓库实现编制。"],
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    make_diagrams()
    paths = [
        requirements_doc(),
        plan_doc(),
        database_doc(),
        architecture_doc(),
        test_doc(),
        meeting_doc(),
        manual_doc(),
        assignments_xlsx(),
    ]
    build_manifest(paths)
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
